import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

p_top = doc.add_paragraph()
r_top = p_top.add_run('붙임 4.  제안 요약서')
r_top.font.size = Pt(16)
r_top.font.bold = True
r_top.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
r_sub = p_top.add_run('  ※ A4 5장 내외')
r_sub.font.size = Pt(10)
r_sub.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [Inches(1.5), Inches(5.3)]
for row in table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

table.rows[0].cells[0].text = '접수번호'
table.rows[0].cells[1].text = '이 칸은 비워두세요.'
table.rows[1].cells[0].text = '팀명'
table.rows[1].cells[1].text = '온중'
table.rows[2].cells[0].text = '제안서 제목'
table.rows[2].cells[1].text = '대구 폭염 대응 실시간 건물 그림자 연산 기반 무더위쉼터 그늘길 안내 및 블록체인 안심 대피 서비스, 「온중(溫證)」'

for r_idx, row in enumerate(table.rows):
    set_cell_background(row.cells[0], 'F3F4F6')
    for c in row.cells:
        set_cell_margins(c, 120, 120, 150, 150)
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.size = Pt(10)
        if c == row.cells[0]:
            p.runs[0].font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

def add_section_header(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x6E, 0x6B)

def add_subsection_header(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def add_body_p(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    r.font.name = '맑은 고딕'
    r.font.size = Pt(9.5)
    return p

add_section_header('[과제요약]  * 500자 이내 작성')
summary_box = doc.add_table(rows=1, cols=1)
summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_box.rows[0].cells[0].width = Inches(6.8)
set_cell_background(summary_box.rows[0].cells[0], 'F9FAFB')
set_cell_margins(summary_box.rows[0].cells[0], 140, 140, 180, 180)
p_sum = summary_box.rows[0].cells[0].paragraphs[0]
p_sum.paragraph_format.line_spacing = 1.3
r_sum = p_sum.add_run('전국 최고 폭염 도시 대구에서 여름철 야외 보행은 온열질환의 주원인입니다. 기존 지도 앱은 직사광선을 고려하지 않고 땡볕 최단거리만 안내하며 무더위쉼터 접근 경로를 제시하지 못합니다.\n본 제안 「온중(溫證)」은 실시간 태양 궤적(고도·방위각)과 대구시 건축물 3D 높이 데이터를 융합해 건물 그림자 영역을 실시간 연산하고 직사광선을 피해 이동하는 \'무더위쉼터 최적 그늘길 내비게이션\'을 제공합니다. 대구 시내 100개 iM뱅크 지점 및 950개 쉼터로의 무계단·완경사 시니어 맞춤형 보행 경로(0.75m/s 기준)를 안내합니다. 기상청 500m 체감온도 기반 폭염 위험 모니터링과 Base Sepolia EAS 블록체인 기반 쉼터 운영상태 및 체크인 신뢰 검증을 결합하여 대구 시민의 보행 안전을 확보하고 iM뱅크의 시니어 포용금융 창구화를 실현하는 생활밀착형 플랫폼입니다. (486자)')
r_sum.font.name = '맑은 고딕'
r_sum.font.size = Pt(9.5)

add_section_header('1) 제안내용의 목적 및 필요성')
add_subsection_header('※ 문제 정의: 기획한 아이디어가 해결하고자 하는 핵심 문제점과 타깃 대상')
add_body_p('• 대구 폭염 속 직사광선 보행 위험: 대구는 분지 지형과 높은 도시 밀도로 여름철 아스팔트 지표 온도가 50~60℃에 달해, 한낮 보행 시 고령자 온열질환 위험이 급증합니다.')
add_body_p('• 기존 지도 앱의 한계: 상용 지도 앱은 땡볕 최단거리만 안내하여 직사광선 대로변으로 유도합니다. (땡볕 100m보다 그늘길 300m가 안전)')
add_body_p('• 보행 약자 특성 및 쉼터 인지 부족: 노인은 보행속도가 느리고(0.75m/s) 계단에 취약하나, 대구 950개 쉼터(iM뱅크 100곳 포함)로의 그늘길 경로가 없습니다.')

add_subsection_header('※ 핵심 솔루션: 개발한 프로토타입이 문제를 해결하는 핵심 메커니즘 요약')
add_body_p('• 실시간 건물 그림자 라우팅: 태양 고도/방위각과 건물 3D 높이를 연산하여 직사광선을 피하는 최적 그늘길 도출.')
add_body_p('• 시니어 무장애 보행: TMAP 보행자 API 연동으로 계단을 100% 배제하고 완경사·휴식 거점 중심 경로 안내.')
add_body_p('• iM뱅크 100개 지점 쉼터화 & 블록체인 검증: 100개 iM뱅크 지점으로 안전 대피를 유도하고 Base Sepolia EAS로 쉼터 체크인 증명.')

add_subsection_header('※ 주요 기능 한눈에 보기: 프로토타입에 구현된 주요 핵심 기능 3~4가지 서술')
add_body_p('• 기능 1. 무더위쉼터 그늘길 내비게이션 (F-04): 실시간 건물 그림자 영역 연산 기반 직사광선 회피 및 계단 없는 길안내.')
add_body_p('• 기능 2. 초국소 체감온도 모니터링 (F-01): 기상청 500m 체감온도 및 폭염 특보 연동을 통한 대피 안내.')
add_body_p('• 기능 3. iM뱅크 지점(100곳) 쉼터 연계 (F-03): 대구 100개 iM뱅크 지점을 쾌적한 금융 쉼터로 우선 추천.')
add_body_p('• 기능 4. Base Sepolia EAS 온체인 증명 (F-05): 쉼터 체크인 및 상태 제보를 HMAC 익명화로 온체인 영구 기록.')

add_section_header('2) 구조 및 기술')
add_subsection_header('※ 시스템 구조도 (이미지화로 대략적인 기술 프로세스 설명)')

arch_box = doc.add_table(rows=1, cols=1)
arch_box.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_box.rows[0].cells[0].width = Inches(6.8)
set_cell_background(arch_box.rows[0].cells[0], 'F3F4F6')
set_cell_margins(arch_box.rows[0].cells[0], 120, 120, 150, 150)
p_arch = arch_box.rows[0].cells[0].paragraphs[0]
p_arch.paragraph_format.line_spacing = 1.2
r_arch = p_arch.add_run(
'''[사용자 화면] Paper 표면(야외 고대비 라이트 테마, 22px 시니어 모드) / Shade 표면(관제판)
       │ (TanStack Start + React 19 SSR)
[연산 엔진] 실시간 그림자 라우팅 파이프라인 (Bun Serverless)
       ├─ 1. SunCalc: 태양 고도(θ) 및 방위각(φ) 실시간 연산
       ├─ 2. 건축물 3D 높이(h) 기반 도로 위 그림자 길이 L = h / tan(θ) 투영
       ├─ 3. Turf.js: 도로망 세그먼트와 그림자 폴리곤 공간 교차 연산
       └─ 4. TMAP 보행자 API: 계단 제외, 노인 보행속도(0.75m/s) 적용 최적 경로
       │
[데이터 & 체인] Supabase PostGIS(950개 쉼터) / 기상청 500m 체감온도 / Base Sepolia EAS'''
)
r_arch.font.name = '맑은 고딕'
r_arch.font.size = Pt(8.5)

add_subsection_header('※ 기술 스택 (Tech Stack Table): 사용한 언어, 프레임워크, 라이브러리, DB, 데이터셋 등 명시')

t_stack = doc.add_table(rows=7, cols=3)
t_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
t_stack.autofit = False
for r in t_stack.rows:
    r.cells[0].width = Inches(1.3)
    r.cells[1].width = Inches(2.2)
    r.cells[2].width = Inches(3.3)

headers = ['영역', '적용 기술', '적용 목적 및 특장점']
for i, h in enumerate(headers):
    t_stack.rows[0].cells[i].text = h
    set_cell_background(t_stack.rows[0].cells[i], 'E5E7EB')
    p = t_stack.rows[0].cells[i].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

data = [
    ('그림자 연산', 'suncalc, @turf/turf', '태양 궤적 및 3D 건물 그림자 폴리곤 연산, 공간 교차 판정'),
    ('지도 / 경로', 'Naver Maps v3, TMAP API', '그늘길 시각화, 계단 제외 및 노인 보행(0.75m/s) 라우팅'),
    ('프론트엔드', 'TanStack Start, React 19', '초고속 SSR 렌더링, 야외용 Paper 표면 및 시니어 모드'),
    ('백엔드 / DB', 'Bun 1.3, Supabase (PostGIS)', '대구 950개 쉼터 반경 검색 및 공간 데이터 고속 인덱싱'),
    ('공공데이터', '기상청 APIHub, 국토부 GIS', '500m 격자 초국소 체감온도 수신, 대구 전역 건축물 높이 데이터'),
    ('블록체인', 'Base Sepolia, EAS SDK', '쉼터 체크인 및 상태 제보의 위변조 방지 온체인 증명서 발급')
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, text in enumerate(row_data):
        c = t_stack.rows[row_idx].cells[col_idx]
        c.text = text
        set_cell_margins(c, 80, 80, 100, 100)
        p = c.paragraphs[0]
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.size = Pt(8.5)
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(c, 'F9FAFB')

add_section_header('3) 제안내용의 차별성')
add_subsection_header('※ 기존 유사 서비스와의 차별점 내용 작성')

t_diff = doc.add_table(rows=6, cols=3)
t_diff.alignment = WD_TABLE_ALIGNMENT.CENTER
t_diff.autofit = False
for r in t_diff.rows:
    r.cells[0].width = Inches(1.5)
    r.cells[1].width = Inches(2.6)
    r.cells[2].width = Inches(2.7)

diff_headers = ['비교 항목', '기존 일반 지도 / 안부전화', '온중 (본 제안 솔루션)']
for i, h in enumerate(diff_headers):
    t_diff.rows[0].cells[i].text = h
    set_cell_background(t_diff.rows[0].cells[i], 'E5E7EB')
    p = t_diff.rows[0].cells[i].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

diff_data = [
    ('경로 기준', '물리적 최단거리 (땡볕 대로변 노출)', '실시간 태양 위치 기반 그늘길 우선 안내'),
    ('보행 약자 맞춤', '성인 기준(1.3m/s), 계단/육교 포함', '노인 속도(0.75m/s), 계단 100% 회피, 완경사'),
    ('목적지 연계', '사용자가 직접 쉼터 검색 필요', '대구 950개 쉼터 및 iM뱅크 100곳 자동 매칭'),
    ('기록 신뢰성', '기록 없음 또는 내부 DB 수기 의존', 'Base Sepolia EAS 블록체인 영구 검증'),
    ('금융권 연계', '단순 지점 위치 검색', 'iM뱅크 100개 지점을 포용금융 쉼터로 연계')
]

for row_idx, row_data in enumerate(diff_data, start=1):
    for col_idx, text in enumerate(row_data):
        c = t_diff.rows[row_idx].cells[col_idx]
        c.text = text
        set_cell_margins(c, 80, 80, 100, 100)
        p = c.paragraphs[0]
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.size = Pt(8.5)
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(c, 'F9FAFB')
        elif col_idx == 2:
            set_cell_background(c, 'F0FDF4')

add_section_header('4) 기대효과')
add_subsection_header('※ 제안한 내용의 기대효과 작성')
add_body_p('• 대구 시민 보행 안전: 한낮 직사광선 노출을 최소화하고 안전한 쉼터로 유도하여 온열질환 사고 및 응급실 내원율 감소.')
add_body_p('• 쉼터 인프라 이용 활성화: 대구 950개 쉼터의 접근성을 높여 공공 복지 인프라 투자 대비 효용 증대.')
add_body_p('• iM뱅크 100개 지점 쉼터 거점화: 대구은행 영업점을 시원한 휴식 공간으로 개방하여 친근한 지역 은행 브랜드 이미지 제고.')
add_body_p('• 시니어 포용금융 창구화: 지점 방문 어르신 대상 시니어 금융 상담(연금, 자산관리) 및 보이스피싱 예방 교육 연계.')
add_body_p('• ESG(Social) 경영 선도: 지역사회 안전망 구축에 직접 기여하여 금융권 ESG 평가 최고 등급 획득.')

add_section_header('5) (자유 기재) 기술 사진, 이미지, 도표 등')
add_subsection_header('※ 제안내용에 대한 기타 추가내용이 있을 경우 작성')

add_body_p('[UI/UX 디자인 시스템]')
add_body_p('• Paper 모바일: 직사광선 아래 시인성을 높인 고대비 라이트 테마 및 22px 시니어 모드 폰트 적용.')
add_body_p('• 그늘 경로 시각화: 건물 그림자 그늘 구간(청록색)과 직사광선 구간(주황색)을 명확히 구분 표시.')

add_body_p('[EAS 온체인 스마트 컨트랙트 데이터 구조 (프라이버시 보호)]')
add_body_p('• HMAC-SHA-256 익명화 기법을 적용하여 개인 식별자 노출 없이 쉼터 체크인 사실만을 Base Sepolia EAS 온체인에 영구 기록 및 공개 검증.')

add_body_p('[사용자 시연 시나리오]')
add_body_p('1. 폭염 발생: 대구 체감온도 38.2℃ 도달.')
add_body_p('2. 쉼터 탐색: 온중 앱 접속 후 클릭 한 번으로 반경 500m 내 iM뱅크 수성동지점 매칭.')
add_body_p('3. 그늘길 도출: 건물 그림자가 드리워진 이면도로 그늘길(계단 없음, 완경사) 안내.')
add_body_p('4. 도착 및 체크인: iM뱅크 지점 도착 후 체크인, Base Sepolia EAS에 증명서 발행 완료.')

output_docx = '붙임4_제안요약서_온중.docx'
doc.save(output_docx)
print(f'Successfully generated {output_docx}')
